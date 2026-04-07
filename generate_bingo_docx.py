import argparse
import json
import random
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Set, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table

PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "LETTER": (21.59, 27.94),
}

NAMED_COLORS = {
    "black": "000000",
    "white": "FFFFFF",
    "grey": "808080",
    "gray": "808080",
    "light-grey": "D3D3D3",
    "light-gray": "D3D3D3",
    "lightgrey": "D3D3D3",
    "lightgray": "D3D3D3",
    "red": "FF0000",
    "green": "008000",
    "blue": "0000FF",
}


def _normalize_hex_color(color_value: str) -> str:
    value = color_value.strip().lower()
    value = NAMED_COLORS.get(value, value)
    value = value[1:] if value.startswith("#") else value
    value = value.lower()
    if len(value) != 6 or any(ch not in "0123456789abcdef" for ch in value):
        raise ValueError(
            "Color values must be a named color (e.g., grey) or 6-digit hex (e.g., #808080)."
        )
    return value.upper()


def _ticket_height_cm(config: Dict[str, Any]) -> float:
    return (
        float(config["round_height_cm"])
        + float(config["header_height_cm"])
        + (float(config["grid_size"]) * float(config["cell_height_cm"]))
        + float(config["footer_height_cm"])
    )


def load_config(config_path: str) -> Dict[str, Any]:
    """Load and validate configuration from JSON or YAML."""
    path = Path(config_path)
    if not path.exists():
        raise FileNotFoundError(f"Config file not found: {config_path}")

    if path.suffix.lower() in {".yaml", ".yml"}:
        try:
            import yaml
        except ImportError as exc:
            raise ImportError(
                "PyYAML is required for YAML configs. Install with: pip install pyyaml"
            ) from exc
        with path.open("r", encoding="utf-8-sig") as file:
            config = yaml.safe_load(file)
    elif path.suffix.lower() == ".json":
        with path.open("r", encoding="utf-8-sig") as file:
            config = json.load(file)
    else:
        raise ValueError("Unsupported config format. Use .json, .yaml, or .yml")

    required_keys = [
        "default_title",
        "rounds",
        "tickets_per_round",
        "grid_size",
        "cell_width_cm",
        "cell_height_cm",
        "tickets_per_row",
        "tickets_per_page",
    ]
    missing = [key for key in required_keys if key not in config]
    if missing:
        raise ValueError(f"Missing required config keys: {', '.join(missing)}")

    config["rounds"] = int(config["rounds"])
    config["tickets_per_round"] = int(config["tickets_per_round"])
    config["grid_size"] = int(config["grid_size"])
    config["cell_width_cm"] = float(config["cell_width_cm"])
    config["cell_height_cm"] = float(config["cell_height_cm"])
    config["tickets_per_row"] = int(config["tickets_per_row"])
    config["tickets_per_page"] = int(config["tickets_per_page"])
    config["page_size"] = str(config.get("page_size", "A4")).upper()
    config["orientation"] = str(config.get("orientation", config.get("page_orientation", "landscape"))).lower()
    config["margin_cm"] = float(config.get("margin_cm", 0.0))
    config["auto_fit_cells"] = bool(config.get("auto_fit_cells", False))
    config["round_height_cm"] = float(config.get("round_height_cm", 0.5))
    config["header_height_cm"] = float(config.get("header_height_cm", 0.5))
    config["footer_height_cm"] = float(config.get("footer_height_cm", 0.3))
    config["text_font"] = str(config.get("text_font", "Calibri"))
    config["text_size"] = int(config.get("text_size", config.get("font_size", 12)))
    config["header_font_size"] = int(config.get("header_font_size", 16))
    config["header_font_style"] = str(config.get("header_font_style", "all-caps")).strip().lower()
    config["header_all_caps"] = config["header_font_style"] in {"all-caps", "all_caps", "caps", "uppercase"}
    config["header_color"] = _normalize_hex_color(str(config.get("header_color", "white")))
    config["header_bg_color"] = _normalize_hex_color(str(config.get("header_bg_color", "grey")))
    config["footer_font_size"] = int(config.get("footer_font_size", 10))
    config["footer_bg_color"] = _normalize_hex_color(str(config.get("footer_bg_color", "white")))
    config["footer_color"] = _normalize_hex_color(str(config.get("footer_color", "black")))
    _grid_border_map = {"thin": 4, "normal": 8, "thick": 16}
    _grid_border_val = str(config.get("grid_border", "normal")).strip().lower()
    config["grid_border_size"] = _grid_border_map.get(_grid_border_val, 8)
    config["free_cell_bg_color"] = _normalize_hex_color(str(config.get("free_cell_bg_color", "white")))
    config["free_cell_text_color"] = _normalize_hex_color(str(config.get("free_cell_text_color", "000000")))

    if config["page_size"] not in PAGE_SIZES_CM:
        raise ValueError(f"page_size must be one of: {', '.join(PAGE_SIZES_CM.keys())}.")
    if config["orientation"] not in {"portrait", "landscape"}:
        raise ValueError("orientation must be either 'portrait' or 'landscape'.")
    if config["margin_cm"] < 0:
        raise ValueError("margin_cm must be >= 0.")
    if config["round_height_cm"] <= 0 or config["header_height_cm"] <= 0 or config["footer_height_cm"] <= 0:
        raise ValueError("round_height_cm, header_height_cm, and footer_height_cm must be > 0.")

    grid_size = config["grid_size"]
    if grid_size <= 0 or grid_size % 2 == 0:
        raise ValueError("grid_size must be a positive odd number (e.g., 7).")

    tickets_per_row = config["tickets_per_row"]
    tickets_per_page = config["tickets_per_page"]
    if tickets_per_row <= 0 or tickets_per_page <= 0:
        raise ValueError("tickets_per_row and tickets_per_page must be positive integers.")
    if tickets_per_page % tickets_per_row != 0:
        raise ValueError("tickets_per_page must be divisible by tickets_per_row.")

    if "tasks" in config:
        config["tasks"] = list(dict.fromkeys(config["tasks"]))
    if "criteria_pool" in config:
        config["criteria_pool"] = list(dict.fromkeys(config["criteria_pool"]))
    config["rows_per_page"] = tickets_per_page // tickets_per_row
    return config


def _apply_page_setup(document: DocxDocument, config: Dict[str, Any]) -> None:
    orientation = config["orientation"]
    margin_cm = float(config["margin_cm"])
    page_width_cm, page_height_cm = PAGE_SIZES_CM[config["page_size"]]

    for section in document.sections:
        section.page_width = Cm(page_width_cm)
        section.page_height = Cm(page_height_cm)

        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width

        margin = Cm(margin_cm)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin


def _set_table_cell_margins(cell: _Cell, top: int = 0, right: int = 0, bottom: int = 0, left: int = 0) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _set_cell_border(cell: _Cell, size: int = 8, color: str = "000000", space: int = 0) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_element = tc_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), str(size))
        edge_element.set(qn("w:space"), str(space))
        edge_element.set(qn("w:color"), color)


def _set_cell_background(cell: _Cell, color_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def _remove_cell_border(cell: _Cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_element = tc_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_element)
        edge_element.set(qn("w:val"), "nil")


def _set_cell_text(
    cell: _Cell,
    text: str,
    font_size: int,
    font_name: str,
    color_hex: str = "000000",
    bold: bool = False,
    all_caps: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.all_caps = all_caps
    run.font.color.rgb = RGBColor.from_string(color_hex)

    # Ensure East Asian font fallback follows selected text font.
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _table_target_add_table(target: Any, rows: int, cols: int) -> Table:
    if hasattr(target, "add_table"):
        try:
            return target.add_table(rows=rows, cols=cols)
        except TypeError:
            return target.add_table(rows=rows, cols=cols, width=Cm(1))
    raise TypeError("Target object cannot host a table.")


def _auto_fit_cells(document: DocxDocument, config: Dict[str, Any]) -> None:
    """Overwrite cell_width_cm and cell_height_cm so tickets fill the printable page exactly."""
    section = document.sections[0]
    emu_to_cm = 360000.0
    page_width = int(section.page_width or 0)
    page_height = int(section.page_height or 0)
    left_margin = int(section.left_margin or 0)
    right_margin = int(section.right_margin or 0)
    top_margin = int(section.top_margin or 0)
    bottom_margin = int(section.bottom_margin or 0)

    printable_width_cm = (page_width - left_margin - right_margin) / emu_to_cm
    printable_height_cm = (page_height - top_margin - bottom_margin) / emu_to_cm

    fixed_ticket_height_cm = (
        float(config["round_height_cm"])
        + float(config["header_height_cm"])
        + float(config["footer_height_cm"])
    )
    available_grid_height_cm = printable_height_cm - (float(config["rows_per_page"]) * fixed_ticket_height_cm)
    if available_grid_height_cm <= 0:
        raise ValueError(
            "Configured round/header/footer heights leave no room for the grid. "
            "Reduce fixed row heights, reduce tickets per page, or change page settings."
        )

    # Divide available space evenly across all ticket columns and grid rows.
    config["cell_width_cm"] = printable_width_cm / (config["tickets_per_row"] * config["grid_size"])
    config["cell_height_cm"] = available_grid_height_cm / (config["rows_per_page"] * config["grid_size"])


def _validate_layout_fits_page(document: DocxDocument, config: Dict[str, Any]) -> None:
    section = document.sections[0]

    emu_to_cm = 360000.0
    page_width = int(section.page_width or 0)
    page_height = int(section.page_height or 0)
    left_margin = int(section.left_margin or 0)
    right_margin = int(section.right_margin or 0)
    top_margin = int(section.top_margin or 0)
    bottom_margin = int(section.bottom_margin or 0)

    printable_width_cm = (page_width - left_margin - right_margin) / emu_to_cm
    printable_height_cm = (page_height - top_margin - bottom_margin) / emu_to_cm

    ticket_width_cm = float(config["grid_size"]) * float(config["cell_width_cm"])
    ticket_height_cm = _ticket_height_cm(config)

    required_width_cm = float(config["tickets_per_row"]) * ticket_width_cm
    required_height_cm = float(config["rows_per_page"]) * ticket_height_cm

    if required_width_cm > printable_width_cm or required_height_cm > printable_height_cm:
        raise ValueError(
            "Configured cell size and page layout do not fit the current page. "
            f"Required area is {required_width_cm:.2f}cm x {required_height_cm:.2f}cm, "
            f"but printable area is {printable_width_cm:.2f}cm x {printable_height_cm:.2f}cm. "
            "Reduce cell dimensions, reduce tickets per page, or change page settings."
        )


def generate_ticket_data(
    round_number: int,
    ticket_number: int,
    config: Dict[str, Any],
    used_signatures: Set[Tuple[str, ...]],
    max_attempts: int = 500,
) -> Dict[str, Any]:
    """Generate one unique ticket payload."""
    grid_size = int(config["grid_size"])
    required_criteria = grid_size * grid_size - 1
    center = grid_size // 2

    criteria_pool = config["criteria_pool"]
    tasks = config["tasks"]

    for _ in range(max_attempts):
        picked = random.sample(criteria_pool, required_criteria)
        grid: List[List[str]] = []
        idx = 0
        for row in range(grid_size):
            row_values: List[str] = []
            for col in range(grid_size):
                if row == center and col == center:
                    row_values.append("⭐")
                else:
                    row_values.append(picked[idx])
                    idx += 1
            grid.append(row_values)

        signature = tuple(value for row_values in grid for value in row_values)
        if signature in used_signatures:
            continue

        used_signatures.add(signature)
        task = random.choice(tasks)
        id_width = max(3, len(str(int(config["tickets_per_round"]))))
        ticket_id = f"R{round_number}-{ticket_number:0{id_width}d}"
        return {
            "round": round_number,
            "ticket_id": ticket_id,
            "title": config["default_title"],
            "grid": grid,
            "task": task,
        }

    raise RuntimeError(
        f"Unable to create a unique ticket after {max_attempts} attempts for round {round_number}."
    )


def create_ticket_table(doc: Any, ticket_data: Dict[str, Any], config: Dict[str, Any]) -> Table:
    """Create one ticket table in the provided document or cell."""
    grid_size = int(config["grid_size"])
    total_rows = 2 + grid_size + 1
    text_size = int(config["text_size"])
    text_font = str(config["text_font"])
    header_font_size = int(config["header_font_size"])
    header_all_caps = bool(config["header_all_caps"])
    header_color = str(config["header_color"])
    header_bg_color = str(config["header_bg_color"])
    footer_font_size = int(config["footer_font_size"])
    footer_bg_color = str(config["footer_bg_color"])
    footer_color = str(config["footer_color"])
    grid_border_size = int(config["grid_border_size"])
    free_cell_bg_color = str(config["free_cell_bg_color"])
    free_cell_text_color = str(config["free_cell_text_color"])

    table = _table_target_add_table(doc, rows=total_rows, cols=grid_size)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    table.rows[0].height = Cm(float(config["round_height_cm"]))
    table.rows[1].height = Cm(float(config["header_height_cm"]))
    for row_index in range(2, 2 + grid_size):
        table.rows[row_index].height = Cm(float(config["cell_height_cm"]))
    table.rows[grid_size + 2].height = Cm(float(config["footer_height_cm"]))

    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(float(config["cell_width_cm"]))
            _set_table_cell_margins(cell, top=0, right=0, bottom=0, left=0)
            _set_cell_border(cell, size=grid_border_size)
            _set_cell_text(cell, "", font_size=text_size, font_name=text_font)

    header_row_1 = table.rows[0].cells
    merged_header_1 = header_row_1[0]
    for col in range(1, grid_size):
        merged_header_1 = merged_header_1.merge(header_row_1[col])
    _set_cell_background(merged_header_1, header_bg_color)
    _set_cell_text(
        merged_header_1,
        f"Round {ticket_data['round']}",
        font_size=header_font_size,
        font_name=text_font,
        color_hex=header_color,
        bold=True,
        all_caps=header_all_caps,
    )

    header_row_2 = table.rows[1].cells
    merged_header_2 = header_row_2[0]
    for col in range(1, grid_size):
        merged_header_2 = merged_header_2.merge(header_row_2[col])
    _set_cell_background(merged_header_2, header_bg_color)
    _set_cell_text(
        merged_header_2,
        ticket_data["title"],
        font_size=header_font_size,
        font_name=text_font,
        color_hex=header_color,
        bold=True,
        all_caps=header_all_caps,
    )

    for row in range(grid_size):
        for col in range(grid_size):
            value = ticket_data["grid"][row][col]
            is_free = value == "⭐"
            cell = table.cell(row + 2, col)
            if is_free:
                _set_cell_background(cell, free_cell_bg_color)
            _set_cell_text(
                cell,
                value,
                font_size=text_size,
                font_name=text_font,
                bold=is_free,
                color_hex=free_cell_text_color if is_free else "000000",
            )

    footer_cells = table.rows[grid_size + 2].cells
    task_cols = grid_size - 2  # leave last 2 cols for ticket ID
    merged_footer = footer_cells[0]
    for col in range(1, task_cols):
        merged_footer = merged_footer.merge(footer_cells[col])
    _set_cell_background(merged_footer, footer_bg_color)
    _set_cell_text(
        merged_footer,
        f"⭐ {ticket_data['task']}",
        font_size=footer_font_size,
        font_name=text_font,
        color_hex=footer_color,
        bold=False,
    )

    ticket_meta = table.cell(grid_size + 2, task_cols).merge(table.cell(grid_size + 2, grid_size - 1))
    _set_cell_background(ticket_meta, footer_bg_color)
    _set_cell_text(
        ticket_meta,
        ticket_data["ticket_id"],
        font_size=footer_font_size,
        font_name=text_font,
        color_hex=footer_color,
        bold=False,
    )

    return table


def _remove_trailing_paragraph(doc: DocxDocument) -> None:
    """Remove the empty paragraph Word auto-inserts after a table."""
    body = doc.element.body
    last = body[-1]
    if last.tag.endswith("}p") and not last.text_content().strip():
        body.remove(last)


def build_page_layout(doc: DocxDocument, tickets: List[Dict[str, Any]], config: Dict[str, Any]) -> None:
    """Build container pages and place one ticket in each container cell."""
    tickets_per_page = int(config["tickets_per_page"])
    tickets_per_row = int(config["tickets_per_row"])
    rows_per_page = int(config["rows_per_page"])
    ticket_width_cm = float(config["grid_size"]) * float(config["cell_width_cm"])
    ticket_height_cm = _ticket_height_cm(config)

    # Remove the default empty paragraph at the start of the document.
    body = doc.element.body
    if body[0].tag.endswith("}p"):
        body.remove(body[0])

    for idx in range(0, len(tickets), tickets_per_page):
        page_tickets = tickets[idx : idx + tickets_per_page]

        container = doc.add_table(rows=rows_per_page, cols=tickets_per_row)
        container.autofit = False
        container.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in container.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Cm(ticket_height_cm)

        for row_cells in container.rows:
            for cell in row_cells.cells:
                cell.width = Cm(ticket_width_cm)
                _set_table_cell_margins(cell, top=0, right=0, bottom=0, left=0)
                _remove_cell_border(cell)
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)

        for ticket_index, ticket in enumerate(page_tickets):
            row = ticket_index // tickets_per_row
            col = ticket_index % tickets_per_row
            cell = container.cell(row, col)
            create_ticket_table(cell, ticket, config)
            # Word requires exactly one <w:p> per cell. Remove all existing ones,
            # then append a single empty <w:p> at the end (after the nested table).
            tc = cell._tc
            for p in list(tc):
                if p.tag == qn("w:p"):
                    tc.remove(p)
            empty_p = OxmlElement("w:p")
            ppr = OxmlElement("w:pPr")
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            ppr.append(spacing)
            empty_p.append(ppr)
            tc.append(empty_p)

        # Remove the empty paragraph Word inserts after the table.
        _remove_trailing_paragraph(doc)

        if idx + tickets_per_page < len(tickets):
            doc.add_page_break()


def _generate_all_tickets(config: Dict[str, Any]) -> Tuple[List[Dict[str, Any]], List[List[Dict[str, Any]]]]:
    all_tickets: List[Dict[str, Any]] = []
    by_round: List[List[Dict[str, Any]]] = []
    used_signatures: Set[Tuple[str, ...]] = set()

    rounds = int(config["rounds"])
    tickets_per_round = int(config["tickets_per_round"])

    for round_number in range(1, rounds + 1):
        round_tickets: List[Dict[str, Any]] = []
        for ticket_number in range(1, tickets_per_round + 1):
            ticket = generate_ticket_data(round_number, ticket_number, config, used_signatures)
            all_tickets.append(ticket)
            round_tickets.append(ticket)
        by_round.append(round_tickets)

    return all_tickets, by_round


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Bingo tickets as a .docx file.")
    parser.add_argument(
        "--grid",
        choices=["3", "5", "7"],
        default=None,
        help="Grid size shortcut: 3, 5, or 7. Loads config_{N}x{N}.json automatically.",
    )
    parser.add_argument("--config", default=None, help="Path to JSON/YAML config file.")
    parser.add_argument("--data", default="data.json", help="Path to shared data file (tasks, criteria_pool).")
    parser.add_argument("--output", default=None, help="Output .docx filename.")
    parser.add_argument(
        "--seed",
        type=int,
        default=None,
        help="Optional random seed for reproducible ticket generation.",
    )
    parser.add_argument(
        "--per-round-files",
        action="store_true",
        help="Also export one .docx file per round.",
    )
    args = parser.parse_args()

    if args.grid and args.config:
        parser.error("Use either --grid or --config, not both.")

    if args.grid:
        config_path = f"config_{args.grid}x{args.grid}.json"
        output = args.output or f"bingo_{args.grid}x{args.grid}.docx"
    elif args.config:
        config_path = args.config
        output = args.output or "bingo_tickets.docx"
    else:
        # Default: use global config.json
        config_path = "config.json"
        output = args.output or "bingo_tickets.docx"

    if args.seed is not None:
        random.seed(args.seed)

    config = load_config(config_path)

    # Merge shared data file if tasks/criteria_pool are missing from config.
    if "tasks" not in config or "criteria_pool" not in config:
        data_path = Path(args.data)
        if not data_path.exists():
            raise FileNotFoundError(
                f"Shared data file not found: {args.data}. "
                "Provide tasks and criteria_pool in the config or in a data file."
            )
        if data_path.suffix.lower() in {".yaml", ".yml"}:
            import yaml
            with data_path.open("r", encoding="utf-8-sig") as f:
                shared = yaml.safe_load(f)
        else:
            with data_path.open("r", encoding="utf-8-sig") as f:
                shared = json.load(f)
        config.setdefault("tasks", shared.get("tasks", []))
        config.setdefault("criteria_pool", shared.get("criteria_pool", []))

    # Re-validate tasks and criteria_pool after merging.
    tasks = list(dict.fromkeys(config.get("tasks", [])))
    if not tasks:
        raise ValueError("tasks must contain at least one item.")
    unique_criteria = list(dict.fromkeys(config.get("criteria_pool", [])))
    grid_size = config["grid_size"]
    required_criteria = grid_size * grid_size - 1
    if len(unique_criteria) < required_criteria:
        raise ValueError(
            f"criteria_pool must contain at least {required_criteria} unique entries "
            f"for a {grid_size}x{grid_size} grid (center is FREE)."
        )
    config["tasks"] = tasks
    config["criteria_pool"] = unique_criteria
    all_tickets, tickets_by_round = _generate_all_tickets(config)

    document = Document()
    _apply_page_setup(document, config)
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    if config["auto_fit_cells"]:
        _auto_fit_cells(document, config)
    _validate_layout_fits_page(document, config)
    build_page_layout(document, all_tickets, config)
    document.save(output)

    if args.per_round_files:
        output_path = Path(output)
        for round_number, round_tickets in enumerate(tickets_by_round, start=1):
            round_doc = Document()
            _apply_page_setup(round_doc, config)
            if config["auto_fit_cells"]:
                _auto_fit_cells(round_doc, config)
            _validate_layout_fits_page(round_doc, config)
            build_page_layout(round_doc, deepcopy(round_tickets), config)
            round_output = output_path.with_name(
                f"{output_path.stem}_round_{round_number}{output_path.suffix}"
            )
            round_doc.save(str(round_output))


if __name__ == "__main__":
    main()
